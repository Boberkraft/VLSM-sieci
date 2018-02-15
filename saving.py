from docx.shared import Cm
from docx import Document
import sys


class FakeNetwork:
    """
    Klasa ukrywająca prawie wszystkie atrybuty obiektu.
    Zwraca jedynie pustego stringa bądź atrybut hosts w wypadku gdy
    zostanie o niego poproszona.
    """
    def __init__(self, hosts):
        self.hosts = hosts

    def __getattribute__(self, item):
        """
        Przysłonięcie metody do pobierania atrybutów tak aby wszystkie prócz hosts
        zwracały pustego stringa.

        :param item: parametr
        :return: Pusty string bądź liczbę hostów.
        """
        if item == 'hosts':
            # zwróć się do rodzica o zasłoniętą metodę
            return super(FakeNetwork, self).__getattribute__('hosts')
        else:
            # zwróć pustego stringa
            return ''


def generate_raport(network, subnets, path, answers=True, fullpage=False):
    """
    Generuje raport na podstawie dostarczonych parametrów.
    :param network: BinNetwork służący jako adres ip w poleceniu
    :param subnets: Lista obiektów Binetwork, które zawierają już wszystkie parametry
    :param answers: pokaż odpowiedzi?
    :param fullpage: Wydrukować ilość sprawdzianów tak aby zajeły one całą stronę?
    """
    subnets_list = ', '.join(str(subnet.hosts) for subnet in subnets)
    number_of_repeats = [0, 5, 5, 4, 4, 4, 3, 3, 3, 2, 2, 2, 2, 2, 2, 2]
    rows = len(subnets)

    if answers is False:
        # jeżeli odpowiedzi nie mają być wyświetlane, to podmień podsieci
        # na specjalnych tajnych agentów
        for i in range(len(subnets)):
            subnets[i] = FakeNetwork(subnets[i].hosts)

    document = Document()
    sections = document.sections
    for section in sections:
        # marginesy
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    columns = "hosts allocated_size network_with_mask broadcast first_address last_address".split()
    # w jakiej kolejności mają pojawiać się parametry sieci
    try:
        times = number_of_repeats[rows] if fullpage else 1
    except:
        times = 1
    for _ in range(times):
        p = document.add_paragraph('Imię i nazwisko: ')

        document.add_paragraph("Podziel sieć o adresie " + str(network) + "/" + str(
            network.mask) + " na podsieci po " + subnets_list + " hostów.")

        # tworzenie tabeli i jej nagłówków
        table = document.add_table(rows=len(subnets) + 1, cols=6)
        header = table.rows[0].cells
        header[0].text = "Hosty"
        header[1].text = "Zaalokowane hosty"
        header[2].text = "Adres podsieci"
        header[3].text = "Broadcast"
        header[4].text = "Pierwszy wolny adres"
        header[5].text = "Ostatni wolny adres"
        table.style = 'TableGrid'

        # wypełnij komórki tabeli
        for network_id, row in enumerate(table.rows[1:]):
            for id, col in enumerate(row.cells):
                col.text = str(getattr(subnets[network_id], columns[id]))
        try:
            # zapisz dokument
            document.save(path)
        except:
            print('Wystąpił problem z zapisem', file=sys.stderr)
        if _ != times - 1:
            document.add_paragraph('')
            document.add_paragraph('')
